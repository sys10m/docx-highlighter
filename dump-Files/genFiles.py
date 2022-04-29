# package to install python-docx
import docx

doc = docx.Document(input())

highlightCount = 0
keyword = input()

for currentParagraph in doc.paragraphs:
    if keyword in currentParagraph.text:
      initRunSize = len(currentParagraph.runs)
      for i in range(initRunSize):
        if keyword in currentParagraph.runs[i].text:
          splitedRun = currentParagraph.runs[i].text.split(keyword)
          for x in range(len(splitedRun)):
            addedSplitedRun = currentParagraph.add_run(splitedRun[x])
            copyRunFormat(addedSplitedRun, currentParagraph.runs[i])
            if x < (len(splitedRun) - 1):
              keywordRun = currentParagraph.add_run(keyword)
              copyRunFormat(keywordRun, currentParagraph.runs[i])
              keywordRun.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
              highlightCount += 1
          currentParagraph.runs[i].clear()
        else:
          newRun = currentParagraph.add_run(currentParagraph.runs[i].text)
          copyRunFormat(newRun, currentParagraph.runs[i])
          currentParagraph.runs[i].clear()

      # clear the first set of runs which is already appended at the back
      for clear in range(initRunSize):
        if currentParagraph.runs[clear].text == "":
          r = currentParagraph.runs[clear]._r
          r.getparent().remove(r)

print(highlightCount)
doc.save(input())

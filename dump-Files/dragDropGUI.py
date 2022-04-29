from tkinter import *
from tkinterdnd2 import *
import docx

filePathList = []

def copyRunFormat(newRun, OriginalRun):
    newRun.bold = OriginalRun.bold
    newRun.italic = OriginalRun.italic
    newRun.style = OriginalRun.style
    newRun.underline = OriginalRun.underline
    newRun.font.color.rgb = OriginalRun.font.color.rgb
    newRun.font.color.theme_color = OriginalRun.font.color.theme_color
    newRun.font.all_caps = OriginalRun.font.all_caps
    newRun.font.bold = OriginalRun.font.bold
    newRun.font.complex_script = OriginalRun.font.complex_script
    newRun.font.cs_bold = OriginalRun.font.cs_bold
    newRun.font.cs_italic = OriginalRun.font.cs_italic
    newRun.font.double_strike = OriginalRun.font.double_strike
    newRun.font.emboss = OriginalRun.font.emboss
    newRun.font.hidden = OriginalRun.font.hidden
    newRun.font.highlight_color = OriginalRun.font.highlight_color
    newRun.font.imprint = OriginalRun.font.imprint
    newRun.font.italic = OriginalRun.font.italic
    newRun.font.math = OriginalRun.font.math
    newRun.font.name = newRun.font.name
    newRun.font.no_proof = OriginalRun.font.no_proof
    newRun.font.outline = OriginalRun.font.outline
    newRun.font.rtl = OriginalRun.font.rtl
    newRun.font.shadow = OriginalRun.font.shadow
    newRun.font.size = OriginalRun.font.size
    newRun.font.small_caps = OriginalRun.font.small_caps
    newRun.font.snap_to_grid = OriginalRun.font.snap_to_grid
    newRun.font.spec_vanish = OriginalRun.font.spec_vanish
    newRun.font.strike = OriginalRun.font.strike
    newRun.font.subscript = OriginalRun.font.subscript
    newRun.font.superscript = OriginalRun.font.superscript
    newRun.font.underline = OriginalRun.font.underline
    newRun.font.web_hidden = OriginalRun.font.web_hidden

def addto_listbox(event):
    fileList.insert("end", event.data)
    filePathList.append(event.data)
    filePathList[-1] = filePathList[-1].replace("{", "")
    filePathList[-1] = filePathList[-1].replace("}", "")
    
def clickButton():
    for file in filePathList:
        generateFiles(file)
    fileList.delete(0, END)
    filePathList.clear()

def generateFiles(OriginalFilePath):
    doc = docx.Document(OriginalFilePath)

    #highlightCount = 0
    keyword = textBox.get()
    
    for currentParagraph in doc.paragraphs:
        if keyword in currentParagraph.text:
            initRunSize = len(currentParagraph.runs)
            for i in range(initRunSize):
                if keyword in currentParagraph.runs[i].text:
                    splitedRun = currentParagraph.runs[i].text.split(keyword)
                    for x in range(len(splitedRun)):
                        addedSplitedRun = currentParagraph.add_run(splitedRun[x])
                        copyRunFormat(addedSplitedRun, currentParagraph.runs[i])
                        if x < (len(splitedRun) - 1):
                            keywordRun = currentParagraph.add_run(keyword)
                            copyRunFormat(keywordRun, currentParagraph.runs[i])
                            keywordRun.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
                            #highlightCount += 1
                    currentParagraph.runs[i].clear()
                else:
                    newRun = currentParagraph.add_run(currentParagraph.runs[i].text)
                    copyRunFormat(newRun, currentParagraph.runs[i])
                    currentParagraph.runs[i].clear()

            # clear the first set of runs which is already appended at the back
            for clear in range(initRunSize):
                if currentParagraph.runs[clear].text == "":
                    r = currentParagraph.runs[clear]._r
                    r.getparent().remove(r)

    #print(highlightCount)
    doc.save(OriginalFilePath.replace(".docx", "-%s.docx" % keyword))

mainWindow = TkinterDnD.Tk()
mainWindow.title('Highlight File')
mainWindow.geometry('400x500')
mainWindow.config(bg='#C0C0C0')

frame = Frame(mainWindow)
frame.grid()
frame.pack()

fileList = Listbox(
    frame, 
    width=48,
    height=15,
    selectmode=SINGLE,
    )

fileList.pack(fill=X, side=TOP)
fileList.drop_target_register(DND_FILES)
fileList.dnd_bind('<<Drop>>', addto_listbox)

scroll = Scrollbar(
    frame,
    orient= HORIZONTAL
    )
scroll.pack(side = BOTTOM, fill=X)

textBox = Entry(mainWindow)
textBox.pack()

genButton = Button(mainWindow, text = "Generate highlighted file(s)", command = clickButton)
genButton.pack( side = BOTTOM)


fileList.configure(xscrollcommand=scroll.set)
scroll.config(command=fileList.xview)


mainWindow.mainloop()

